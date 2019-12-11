import os
import docx

x = []
techSummary = {}
fileNames = []
for root, dirs, files in os.walk("."):
    for filename in files:
        if filename.endswith("Resume.docx"):
            fileNames.append(filename)

for resume in fileNames:
    doc = docx.Document(resume)
    for para_index in range(len(doc.paragraphs)):
        if doc.paragraphs[para_index].text.lower() == 'technology summary' or \
                doc.paragraphs[para_index].text.lower() == 'skills':
            var_string = ''
            for string in doc.paragraphs[para_index + 1].runs:
                var_string += string.text
            techSummary[resume] = var_string

for value in techSummary.values():
    data = value.split()
    x.append(data)

print(x)
