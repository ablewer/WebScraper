import docx


def get_text(file_name):
    doc = docx.Document(file_name)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    return '\n'.join(full_text)


print(get_text('RinglerShawn_Resume4.docx'))
